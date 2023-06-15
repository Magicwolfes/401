#!/usr/bin/env python3

import datetime
from docx import Document

Name = input("Please write your First and Last name: ")
today = datetime.datetime.now().date()
Job = input("Please write the Job title you want: ")
Recipient = input("Please write your contact's name: ")
Company = input("Please write the Company's name: ")


address = input("Home Address: ")
State = input("City, State, Zip: ")
email = input("Email: ")
phone = input("Phone number: ")


paragraph1 = input("Paste paragraph one: ")
paragraph2 = input("Paste paragraph one: ")
paragraph3 = input("Paste paragraph one: ")
paragraph4 = input("Paste paragraph one: ")

output_file = "cover_letter.docx"

# Create a new Word document
document = Document()

document.add_paragraph(Name)
document.add_paragraph(str(today))
document.add_paragraph(address)
document.add_paragraph(State)
document.add_paragraph(email)
document.add_paragraph(phone)
document.add_paragraph(Company)
document.add_paragraph()
document.add_paragraph(f"Dear {Recipient},")
document.add_paragraph(paragraph1)
document.add_paragraph(paragraph2)
document.add_paragraph(paragraph3)
document.add_paragraph(paragraph4)
document.add_paragraph("Sincerely, ")
document.add_paragraph(Name)

# Count the number of words in the cover letter
word_count = sum(len(para.text.split()) for para in document.paragraphs)

# Print the word count in the terminal
print("Word count:", word_count)

document.save(output_file)